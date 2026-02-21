"""
Manuscript assembler: reads DB for all phase data and builds a structured DOCX.

Returns base64-encoded DOCX bytes + metadata.
"""
from __future__ import annotations

import base64
from io import BytesIO
from typing import Any, Dict, List, Optional, Tuple
import uuid

from docx import Document
from docx.shared import Inches, Pt
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.grade_assessment import GradeAssessment
from app.models.phase_result import PhaseResult
from app.models.protocol_version import ProtocolVersion
from app.models.screening_decision import ScreeningDecision
from app.models.search_query import SearchQuery
from app.schemas.manuscript import ManuscriptRequest, ManuscriptResponse


PLACEHOLDER = "[PLACEHOLDER — data not yet available for this section]"


async def _fetch_phase_data(
    db: AsyncSession, review_id: uuid.UUID
) -> Dict[str, Any]:
    """Fetch all relevant DB data for a review."""
    data: Dict[str, Any] = {}

    # Phase 1 — latest protocol version
    pv_result = await db.execute(
        select(ProtocolVersion)
        .where(ProtocolVersion.review_id == review_id)
        .order_by(ProtocolVersion.version.desc())
        .limit(1)
    )
    data["protocol"] = pv_result.scalar_one_or_none()

    # Phase 2 — latest search query
    sq_result = await db.execute(
        select(SearchQuery)
        .where(SearchQuery.review_id == review_id)
        .order_by(SearchQuery.created_at.desc())
        .limit(1)
    )
    data["search"] = sq_result.scalar_one_or_none()

    # Phase 3/4 — screening counts
    counts_result = await db.execute(
        select(ScreeningDecision.final_label, func.count(ScreeningDecision.id))
        .where(ScreeningDecision.review_id == review_id)
        .group_by(ScreeningDecision.final_label)
    )
    data["screening_counts"] = dict(counts_result.all())

    # Phases 5–9 — stub results
    stubs_result = await db.execute(
        select(PhaseResult)
        .where(
            PhaseResult.review_id == review_id,
            PhaseResult.phase_number.in_([5, 6, 7, 8, 9]),
        )
    )
    stubs = stubs_result.scalars().all()
    data["phase_stubs"] = {pr.phase_number: pr for pr in stubs}

    # Phase 10 — GRADE assessments
    grade_result = await db.execute(
        select(GradeAssessment).where(GradeAssessment.review_id == review_id)
    )
    data["grade_assessments"] = grade_result.scalars().all()

    return data


def _add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def _add_body(doc: Document, text: str) -> None:
    para = doc.add_paragraph(text)
    para.style.font.size = Pt(11)


def _build_sections(
    doc: Document,
    review_id: uuid.UUID,
    title: Optional[str],
    data: Dict[str, Any],
    methods_narrative: Optional[str],
    results_narrative: Optional[str],
) -> Tuple[List[str], List[str]]:
    """Assemble all sections. Returns (sections_included, missing_phase_data)."""
    sections_included: List[str] = []
    missing_phase_data: List[str] = []

    protocol = data["protocol"]
    search = data["search"]
    screening_counts = data["screening_counts"]
    phase_stubs = data["phase_stubs"]
    grade_assessments = data["grade_assessments"]

    # --- Title ---
    doc_title = title or (
        protocol.research_question[:100] if protocol else "Systematic Review and Meta-Analysis"
    )
    doc.add_heading(doc_title, level=0)
    sections_included.append("Title")

    # --- Abstract ---
    _add_heading(doc, "Abstract", 1)
    if protocol and search:
        pico = protocol.pico_schema or {}
        _add_body(
            doc,
            f"Background: {pico.get('population', PLACEHOLDER)}. "
            f"This review examined the effects of {pico.get('intervention', PLACEHOLDER)} "
            f"compared to {pico.get('comparator', PLACEHOLDER)}.",
        )
        total_screened = sum(screening_counts.values()) if screening_counts else 0
        included = screening_counts.get("include", 0)
        _add_body(
            doc,
            f"Methods: A systematic search was conducted across multiple databases. "
            f"{total_screened} records were screened; {included} studies were included.",
        )
        sections_included.append("Abstract")
    else:
        _add_body(doc, PLACEHOLDER)
        missing_phase_data.append("Abstract (requires Phase 1 + Phase 2 data)")

    # --- Introduction ---
    _add_heading(doc, "Introduction", 1)
    if protocol:
        pico = protocol.pico_schema or {}
        _add_body(
            doc,
            f"This systematic review and meta-analysis addresses the question: "
            f"{protocol.research_question}. "
            f"The population of interest is {pico.get('population', '[population]')}, "
            f"and the primary intervention is {pico.get('intervention', '[intervention]')}.",
        )
        sections_included.append("Introduction")
    else:
        _add_body(doc, PLACEHOLDER)
        missing_phase_data.append("Introduction (requires Phase 1 protocol)")

    # --- Methods ---
    _add_heading(doc, "Methods", 1)

    _add_heading(doc, "Registration", 2)
    _add_body(doc, "This review was registered prospectively on PROSPERO (registration number: [PROSPERO ID]).")
    sections_included.append("Methods: Registration")

    _add_heading(doc, "Eligibility Criteria", 2)
    if protocol:
        pico = protocol.pico_schema or {}
        _add_body(
            doc,
            f"Studies were eligible if they enrolled {pico.get('population', '[population]')}, "
            f"evaluated {pico.get('intervention', '[intervention]')} versus "
            f"{pico.get('comparator', '[comparator]')}, and reported at least one of the following outcomes: "
            f"{', '.join(pico.get('outcomes', ['[outcomes]']))}. "
            f"Eligible study designs included: {', '.join(pico.get('study_designs', ['randomised controlled trial']))}.",
        )
        sections_included.append("Methods: Eligibility")
    else:
        _add_body(doc, PLACEHOLDER)
        missing_phase_data.append("Eligibility criteria (requires Phase 1 PICO)")

    _add_heading(doc, "Information Sources", 2)
    if search:
        _add_body(doc, f"Searches were conducted in {search.database or 'multiple databases'} "
                  f"on {str(search.created_at.date()) if search.created_at else '[date]'}.")
        sections_included.append("Methods: Information sources")
    else:
        _add_body(doc, PLACEHOLDER)
        missing_phase_data.append("Search databases (requires Phase 2 data)")

    _add_heading(doc, "Search Strategy", 2)
    if search and search.search_string:
        _add_body(doc, f"The following search string was used: {search.search_string}")
        sections_included.append("Methods: Search strategy")
    else:
        _add_body(doc, PLACEHOLDER)
        missing_phase_data.append("Search strategy (requires Phase 2 data)")

    _add_heading(doc, "Study Selection", 2)
    if methods_narrative:
        _add_body(doc, methods_narrative)
        sections_included.append("Methods: Study selection (Claude narrative)")
    else:
        _add_body(
            doc,
            "Two independent reviewers screened titles and abstracts, followed by full-text review of eligible records. "
            "Disagreements were resolved by consensus.",
        )
        sections_included.append("Methods: Study selection")

    _add_heading(doc, "Data Extraction", 2)
    if 6 in phase_stubs:
        _add_body(doc, "Data were extracted by two independent reviewers using a pre-specified form.")
        sections_included.append("Methods: Data extraction")
    else:
        _add_body(doc, PLACEHOLDER)
        missing_phase_data.append("Data extraction (requires Phase 6 data)")

    _add_heading(doc, "Risk of Bias Assessment", 2)
    if 7 in phase_stubs:
        rob_data = phase_stubs[7].result_data or {}
        overall = rob_data.get("overall_rob", "not specified")
        _add_body(doc, f"Risk of bias was assessed using [tool]. Overall risk of bias was rated as '{overall}'.")
        sections_included.append("Methods: Risk of bias")
    else:
        _add_body(doc, PLACEHOLDER)
        missing_phase_data.append("Risk of bias assessment (requires Phase 7 data)")

    _add_heading(doc, "Statistical Analysis", 2)
    if 8 in phase_stubs:
        meta_data = phase_stubs[8].result_data or {}
        i2 = meta_data.get("i2", "[I²]")
        _add_body(
            doc,
            f"A random-effects meta-analysis was performed. Heterogeneity was quantified using I² ({i2}%). "
            "Effect estimates are presented with 95% confidence intervals.",
        )
        sections_included.append("Methods: Statistical analysis")
    else:
        _add_body(doc, PLACEHOLDER)
        missing_phase_data.append("Statistical methods (requires Phase 8 data)")

    # --- Results ---
    _add_heading(doc, "Results", 1)

    _add_heading(doc, "Study Selection", 2)
    if screening_counts:
        total = sum(screening_counts.values())
        included = screening_counts.get("include", 0)
        excluded = screening_counts.get("exclude", 0)
        uncertain = screening_counts.get("uncertain", 0)
        body = (
            results_narrative
            if results_narrative
            else (
                f"A total of {total} records were screened at title/abstract stage. "
                f"{included} records were included, {excluded} excluded, and {uncertain} were uncertain."
            )
        )
        _add_body(doc, body)
        sections_included.append("Results: Study selection")
    else:
        _add_body(doc, PLACEHOLDER)
        missing_phase_data.append("Study selection results (requires Phase 3 screening data)")

    _add_heading(doc, "Study Characteristics", 2)
    if 6 in phase_stubs:
        _add_body(doc, "Characteristics of included studies are presented in Table 1.")
        sections_included.append("Results: Study characteristics")
    else:
        _add_body(doc, PLACEHOLDER)
        missing_phase_data.append("Study characteristics (requires Phase 6 extraction)")

    _add_heading(doc, "Risk of Bias", 2)
    if 7 in phase_stubs:
        rob_data = phase_stubs[7].result_data or {}
        _add_body(doc, f"Risk of bias results are summarised in Figure 1. Overall: {rob_data.get('overall_rob', '[see figure]')}.")
        sections_included.append("Results: Risk of bias")
    else:
        _add_body(doc, PLACEHOLDER)
        missing_phase_data.append("Risk of bias results (requires Phase 7 data)")

    _add_heading(doc, "Effects of Interventions", 2)
    if 8 in phase_stubs:
        meta_data = phase_stubs[8].result_data or {}
        rr = meta_data.get("pooled_rr", "[effect estimate]")
        _add_body(doc, f"Meta-analysis yielded a pooled effect of RR={rr} (95% CI: [lower–upper]).")
        sections_included.append("Results: Effects of interventions")
    else:
        _add_body(doc, PLACEHOLDER)
        missing_phase_data.append("Effect estimates (requires Phase 8 meta-analysis)")

    _add_heading(doc, "Publication Bias", 2)
    if 9 in phase_stubs:
        pb_data = phase_stubs[9].result_data or {}
        p = pb_data.get("egger_pval", "[p-value]")
        _add_body(doc, f"Funnel plot asymmetry was assessed using Egger's test (p={p}).")
        sections_included.append("Results: Publication bias")
    else:
        _add_body(doc, PLACEHOLDER)
        missing_phase_data.append("Publication bias (requires Phase 9 data)")

    _add_heading(doc, "Certainty of Evidence (GRADE)", 2)
    if grade_assessments:
        for ga in grade_assessments:
            _add_body(
                doc,
                f"For the outcome '{ga.outcome_name}', the certainty of evidence was rated as "
                f"{ga.certainty} ({ga.downgrade_count} downgrade(s)).",
            )
        sections_included.append("Results: GRADE certainty")
    else:
        _add_body(doc, PLACEHOLDER)
        missing_phase_data.append("GRADE certainty (requires Phase 10 data)")

    # --- Discussion ---
    _add_heading(doc, "Discussion", 1)
    _add_body(
        doc,
        "This systematic review synthesised evidence on the effectiveness of the intervention. "
        "[Discussion to be completed by authors].",
    )
    sections_included.append("Discussion")

    # --- Tables ---
    _add_heading(doc, "Tables", 1)
    _add_body(doc, "Table 1: Characteristics of included studies. [See attached SoF table for summary of findings].")
    sections_included.append("Tables")

    # --- Appendices ---
    _add_heading(doc, "Appendices", 1)
    _add_body(doc, "Appendix A: Full search strategies.\nAppendix B: PRISMA 2020 checklist.")
    sections_included.append("Appendices")

    return sections_included, missing_phase_data


async def build_manuscript(
    db: AsyncSession,
    request: ManuscriptRequest,
) -> ManuscriptResponse:
    data = await _fetch_phase_data(db, request.review_id)

    methods_narrative: Optional[str] = None
    results_narrative: Optional[str] = None

    if request.use_claude_narratives:
        from app.services import claude as claude_svc

        protocol = data["protocol"]
        search = data["search"]
        grade_assessments = data["grade_assessments"]

        if protocol and search:
            pico = protocol.pico_schema or {}
            try:
                methods_narrative = await claude_svc.generate_methods_narrative(
                    pico=pico,
                    search_string=search.search_string or "",
                )
            except Exception:
                methods_narrative = None

        if data["screening_counts"] or grade_assessments:
            grade_list = [
                {"outcome": ga.outcome_name, "certainty": ga.certainty}
                for ga in grade_assessments
            ]
            try:
                results_narrative = await claude_svc.generate_results_narrative(
                    screening_counts=data["screening_counts"],
                    grade_outcomes=grade_list,
                )
            except Exception:
                results_narrative = None

    doc = Document()

    # Narrow margins
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    sections_included, missing_phase_data = _build_sections(
        doc=doc,
        review_id=request.review_id,
        title=request.title,
        data=data,
        methods_narrative=methods_narrative,
        results_narrative=results_narrative,
    )

    # Count words
    full_text = " ".join(
        para.text for para in doc.paragraphs if para.text.strip()
    )
    word_count = len(full_text.split())

    buf = BytesIO()
    doc.save(buf)
    docx_b64 = base64.b64encode(buf.getvalue()).decode()

    return ManuscriptResponse(
        docx_b64=docx_b64,
        sections_included=sections_included,
        missing_phase_data=missing_phase_data,
        word_count=word_count,
    )

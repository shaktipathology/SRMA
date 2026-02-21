"""
Summary of Findings (SoF) table generator using python-docx.

Returns base64-encoded DOCX bytes. No filesystem I/O.
"""
from __future__ import annotations

import base64
from io import BytesIO
from typing import List

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.schemas.sof import SofOutcome, SofRequest


# GRADE certainty → background colour (hex without #)
CERTAINTY_COLOURS = {
    "high": "6FAE3C",
    "moderate": "F7D002",
    "low": "F7941D",
    "very_low": "C8202E",
}

HEADERS = [
    "Outcomes",
    "N participants\n(studies)",
    "Effect measure",
    "Relative effect\n(95% CI)",
    "Absolute effect",
    "Certainty of\nevidence",
    "Importance",
    "Footnotes",
]


def _set_cell_bg(cell, hex_colour: str) -> None:
    """Apply a background fill colour to a table cell via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tcPr.append(shd)


def _bold_cell(cell) -> None:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True


def _set_col_widths(table, widths_inches: List[float]) -> None:
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.width = Inches(widths_inches[j])


def generate_sof_docx(request: SofRequest) -> bytes:
    """Build a SoF DOCX and return raw bytes."""
    doc = Document()

    # Narrow margins
    section = doc.sections[0]
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Title heading
    title_text = request.title or "Summary of Findings Table"
    doc.add_heading(title_text, level=1)

    # Header paragraph
    header_para = doc.add_paragraph()
    header_para.add_run("Population: ").bold = True
    header_para.add_run(request.population + "\n")
    header_para.add_run("Intervention: ").bold = True
    header_para.add_run(request.intervention + "\n")
    header_para.add_run("Comparator: ").bold = True
    header_para.add_run(request.comparator)

    doc.add_paragraph()  # spacer

    # Create 8-column table
    n_rows = 1 + len(request.outcomes)  # header + data rows
    table = doc.add_table(rows=n_rows, cols=8)
    table.style = "Table Grid"

    # Column widths (total ~9.5 inches within narrow margins)
    col_widths = [1.8, 0.9, 0.9, 1.1, 1.1, 1.0, 0.85, 1.35]

    # Header row
    header_row = table.rows[0]
    for j, header_text in enumerate(HEADERS):
        cell = header_row.cells[j]
        cell.text = header_text
        _set_cell_bg(cell, "D9D9D9")  # grey shading
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(9)

    # Data rows
    collected_footnotes: List[str] = []
    footnote_map: dict = {}

    for outcome in request.outcomes:
        for fn in outcome.footnotes:
            if fn not in footnote_map:
                footnote_map[fn] = len(collected_footnotes) + 1
                collected_footnotes.append(fn)

    for row_idx, outcome in enumerate(request.outcomes, start=1):
        row = table.rows[row_idx]

        # Build footnote refs
        fn_refs = ",".join(str(footnote_map[fn]) for fn in outcome.footnotes if fn in footnote_map)
        fn_str = f" [{fn_refs}]" if fn_refs else ""

        relative_effect = f"{outcome.effect_size:.2f} ({outcome.ci_lower:.2f}–{outcome.ci_upper:.2f})"

        values = [
            outcome.outcome_name,
            f"{outcome.n_participants:,}\n({outcome.n_studies} studies)",
            outcome.effect_measure,
            relative_effect,
            "—",  # absolute effect placeholder
            outcome.certainty.replace("_", " ").title(),
            outcome.importance.replace("_", " ").title(),
            fn_str,
        ]

        for j, val in enumerate(values):
            cell = row.cells[j]
            cell.text = val
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(9)

            # Apply certainty colour to the certainty column (index 5)
            if j == 5:
                colour_key = outcome.certainty.lower().replace(" ", "_")
                colour = CERTAINTY_COLOURS.get(colour_key, "FFFFFF")
                _set_cell_bg(cell, colour)
                # White text for darker backgrounds
                if colour_key in ("low", "very_low"):
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Apply column widths
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            cell.width = Inches(col_widths[j])

    # Footnote legend
    if collected_footnotes:
        doc.add_paragraph()
        fn_para = doc.add_paragraph()
        fn_para.add_run("Footnotes:").bold = True
        for idx, fn_text in enumerate(collected_footnotes, start=1):
            doc.add_paragraph(f"[{idx}] {fn_text}", style="List Number")

    # Serialize to bytes
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_sof_b64(request: SofRequest) -> str:
    """Return base64-encoded DOCX string."""
    return base64.b64encode(generate_sof_docx(request)).decode()
